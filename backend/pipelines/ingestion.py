# ─────────────────────────────────────────────────────────────
# FILE: clinicaliq/backend/pipelines/ingestion.py
# ─────────────────────────────────────────────────────────────

import re
import logging
from typing import Optional
import pdfplumber
import docx
from io import BytesIO
from langchain.text_splitter import RecursiveCharacterTextSplitter

from pipelines.ner import extract_entities
from utils.embedder import get_embeddings
from services.chroma import get_chroma_collection
from services.supabase_client import supabase

logger = logging.getLogger(__name__)

# ── Text Splitter — 512 tokens, 50 overlap ─────────────────────
splitter = RecursiveCharacterTextSplitter(
    chunk_size=512,
    chunk_overlap=50,
    separators=["\n\n", "\n", ".", " "]
)

# ── Main Ingestion Function ────────────────────────────────────
async def ingest_document(
    file_bytes: bytes,
    filename: str,
    patient_id: str,
    category: str,
    doc_type: str,
    uploaded_by: str
) -> dict:
    """
    Full ingestion pipeline:
    parse → clean → chunk → NER → embed → store in ChromaDB + Supabase
    """
    parse_error = False
    chunk_count = 0
    entities_found = {}

    try:
        # ── Step 1: Parse ──────────────────────────────────────
        raw_text = parse_file(file_bytes, filename)

        if not raw_text or len(raw_text.strip()) == 0:
            logger.warning(f"Empty or unreadable file: {filename}")
            parse_error = True
            raw_text = ""

    except Exception as e:
        logger.error(f"Parse error for {filename}: {e}")
        parse_error = True
        raw_text = ""

    try:
        # ── Step 2: Clean ──────────────────────────────────────
        cleaned_text = clean_text(raw_text)

        # ── Step 3: Chunk ──────────────────────────────────────
        chunks = splitter.split_text(cleaned_text) if cleaned_text else []
        chunk_count = len(chunks)

        # ── Step 4: NER on each chunk ──────────────────────────
        all_entities = []
        for i, chunk in enumerate(chunks):
            entities = extract_entities(chunk, chunk_index=i)
            all_entities.extend(entities)
        entities_found = {"total": len(all_entities), "entities": all_entities}

        # ── Step 5: Embed chunks ───────────────────────────────
        if chunks:
            embeddings = get_embeddings(chunks)

            # ── Step 6: Store in ChromaDB ──────────────────────
            collection = get_chroma_collection()
            ids = [f"{patient_id}_{filename}_{i}" for i in range(len(chunks))]
            metadatas = [
                {
                    "patient_id":      patient_id,
                    "category":        category,
                    "doc_type":        doc_type,
                    "source_filename": filename,
                    "chunk_index":     i
                }
                for i in range(len(chunks))
            ]
            collection.add(
                ids=ids,
                documents=chunks,
                embeddings=embeddings,
                metadatas=metadatas
            )

        # ── Step 7: Store in Supabase documents table ──────────
        result = supabase.table("documents").insert({
            "uploaded_by": uploaded_by,
            "patient_id":  patient_id,
            "filename":    filename,
            "category":    category,
            "doc_type":    doc_type,
            "chunk_count": chunk_count,
            "parse_error": parse_error
        }).execute()

        doc_id = result.data[0]["id"] if result.data else None

        return {
            "doc_id":         doc_id,
            "chunk_count":    chunk_count,
            "entities_found": entities_found,
            "parse_error":    parse_error
        }

    except Exception as e:
        logger.error(f"Ingestion pipeline error for {filename}: {e}")
        # Degrade gracefully — store error record in Supabase
        supabase.table("documents").insert({
            "uploaded_by": uploaded_by,
            "patient_id":  patient_id,
            "filename":    filename,
            "category":    category,
            "doc_type":    doc_type,
            "chunk_count": 0,
            "parse_error": True
        }).execute()
        return {
            "doc_id":         None,
            "chunk_count":    0,
            "entities_found": {},
            "parse_error":    True
        }


# ── File Parser ────────────────────────────────────────────────
def parse_file(file_bytes: bytes, filename: str) -> str:
    """Parse PDF or DOCX and return raw text."""
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        return parse_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return parse_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}")


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    text_parts = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx — paragraphs + tables."""
    doc = docx.Document(BytesIO(file_bytes))
    parts = []
    # Extract paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


# ── Text Cleaner ───────────────────────────────────────────────
def clean_text(text: str) -> str:
    """
    Clean raw extracted text:
    - Remove page numbers
    - Strip excessive whitespace
    - Remove headers/footers patterns
    """
    # Remove page numbers like "Page 1 of 5"
    text = re.sub(r"Page\s+\d+\s+of\s+\d+", "", text, flags=re.IGNORECASE)

    # Remove standalone numbers (page numbers)
    text = re.sub(r"^\s*\d+\s*$", "", text, flags=re.MULTILINE)

    # Collapse multiple newlines
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Collapse multiple spaces
    text = re.sub(r" {2,}", " ", text)

    return text.strip()